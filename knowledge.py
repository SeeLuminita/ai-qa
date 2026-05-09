"""
Day 6: 企业知识库问答系统
- 上传文档构建知识库
- RAG 智能问答
- Streamlit 界面
"""

import os
import streamlit as st
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain_qdrant import QdrantVectorStore
from qdrant_client import QdrantClient
from qdrant_client.http.models import Distance, VectorParams
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.embeddings import Embeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from typing import List
import dashscope 
from dashscope import TextEmbedding 

load_dotenv()

# ============================================
# Qdrant 配置
# ============================================
def get_qdrant_client():
    """获取Qdrant客户端（支持本地和云端）"""
    qdrant_url = os.getenv("QDRANT_URL", "http://localhost:6333")
    qdrant_api_key = os.getenv("QDRANT_API_KEY", None)
    
    if qdrant_api_key:
        return QdrantClient(url=qdrant_url, api_key=qdrant_api_key)
    else:
        return QdrantClient(url=qdrant_url)

# ============================================
# 百炼 Embedding
# ============================================
class BailianEmbeddings(Embeddings):
    """百炼 Embedding API"""
    
    def __init__(self, model: str = "text-embedding-v3"):
        self.model = model
        dashscope.api_key = os.getenv("DASHSCOPE_API_KEY")
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """批量embedding，每次最多10个文本"""
        all_embeddings = []
        batch_size = 10
        
        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]
            result = TextEmbedding.call(model=self.model, input=batch)
            if result.output is None:
                raise ValueError(f"Embedding API返回None: {result}")
            all_embeddings.extend([item['embedding'] for item in result.output['embeddings']])
        
        return all_embeddings
    
    def embed_query(self, text: str) -> List[float]:
        result = TextEmbedding.call(model=self.model, input=[text])
        if result.output is None:
            raise ValueError(f"Embedding API返回None: {result}")
        return result.output['embeddings'][0]['embedding']

# ============================================
# 页面配置
# ============================================
st.set_page_config(page_title="企业知识库问答", page_icon="📚", layout="wide")

st.title("📚 企业知识库问答系统")
st.write("上传文档，构建知识库，智能问答")

# ============================================
# 文档解析函数
# ============================================
def parse_txt(file):
    """解析 TXT 文件"""
    return file.read().decode("utf-8")

def parse_pdf(file):
    """解析 PDF 文件"""
    from pypdf import PdfReader
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

def parse_docx(file):
    """解析 DOCX 文件"""
    from docx import Document
    doc = Document(file)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def parse_excel(file):
    """解析 Excel 文件"""
    import pandas as pd
    df = pd.read_excel(file)
    # 将每一行转换为文本
    text = ""
    for _, row in df.iterrows():
        row_text = " | ".join([str(val) for val in row.values if pd.notna(val)])
        text += row_text + "\n"
    return text

def parse_file(file):
    """根据文件类型解析文件"""
    file_name = file.name.lower()
    
    if file_name.endswith(".txt"):
        return parse_txt(file)
    elif file_name.endswith(".pdf"):
        return parse_pdf(file)
    elif file_name.endswith(".docx") or file_name.endswith(".doc"):
        return parse_docx(file)
    elif file_name.endswith(".xlsx") or file_name.endswith(".xls"):
        return parse_excel(file)
    else:
        raise ValueError(f"不支持的文件类型: {file_name}")

# ============================================
# 侧边栏：知识库管理
# ============================================
with st.sidebar:
    st.header("📁 知识库管理")
    
    # 上传文档 - 支持多种格式
    uploaded_files = st.file_uploader(
        "上传文档",
        type=["txt", "pdf", "docx", "doc", "xlsx", "xls"],
        accept_multiple_files=True,
        help="支持格式：TXT、PDF、DOCX、Excel"
    )
    
    # 构建选项
    build_mode = st.radio(
        "构建方式",
        ["增量添加", "清空重建"],
        help="增量添加：保留已有文档，追加新文档\n清空重建：删除所有旧文档，重新构建"
    )
    
    col1, col2 = st.columns(2)
    with col1:
        build_btn = st.button("构建知识库", type="primary")
    with col2:
        clear_btn = st.button("清空知识库", type="secondary")
    
    # 清空知识库
    if clear_btn:
        try:
            client = get_qdrant_client()
            client.delete_collection(collection_name="knowledge")
            st.success("✅ 知识库已清空")
        except Exception as e:
            st.success(f"✅ 知识库已清空（{e}）")
    
    # 构建知识库
    if build_btn:
        if uploaded_files:
            with st.spinner("正在构建知识库..."):
                # 读取文档
                documents = []
                for file in uploaded_files:
                    try:
                        content = parse_file(file)
                        if content.strip():  # 确保内容不为空
                            documents.append(Document(
                                page_content=content,
                                metadata={"source": file.name}
                            ))
                            st.write(f"✓ 已解析: {file.name}")
                    except Exception as e:
                        st.error(f"解析失败 {file.name}: {e}")
                
                if documents:
                    # 文档分割
                    text_splitter = RecursiveCharacterTextSplitter(
                        chunk_size=500,
                        chunk_overlap=100,
                        separators=["\n\n", "\n", "。", "！", "？", "；", " ", ""]
                    )
                    splits = text_splitter.split_documents(documents)
                    
                    # 创建或更新向量数据库
                    embeddings = BailianEmbeddings()
                    client = get_qdrant_client()
                    qdrant_url = os.getenv("QDRANT_URL", "http://localhost:6333")
                    qdrant_api_key = os.getenv("QDRANT_API_KEY", None)
                    
                    if build_mode == "清空重建":
                        # 删除旧的collection
                        try:
                            client.delete_collection(collection_name="knowledge")
                        except:
                            pass 
                        
                        # 创建新的collection并添加文档
                        if qdrant_api_key:
                            vectorstore = QdrantVectorStore.from_documents(
                                documents=splits,
                                embedding=embeddings,
                                url=qdrant_url,
                                api_key=qdrant_api_key,
                                collection_name="knowledge"
                            )
                        else:
                            vectorstore = QdrantVectorStore.from_documents(
                                documents=splits,
                                embedding=embeddings,
                                url=qdrant_url,
                                collection_name="knowledge"
                            )
                        st.success(f"✅ 已导入 {len(documents)} 个文档，分割为 {len(splits)} 个片段（已清空旧数据）")
                    else:
                        # 增量添加
                        # 检查collection是否存在
                        collections = client.get_collections().collections
                        collection_names = [c.name for c in collections]
                        
                        if "knowledge" in collection_names:
                            # 加载已有的
                            vectorstore = QdrantVectorStore(
                                client=client,
                                collection_name="knowledge",
                                embedding=embeddings
                            )
                            # 添加新文档
                            vectorstore.add_documents(splits)
                            st.success(f"✅ 已追加 {len(documents)} 个文档，新增 {len(splits)} 个片段")
                        else:
                            # 第一次创建
                            if qdrant_api_key:
                                vectorstore = QdrantVectorStore.from_documents(
                                    documents=splits,
                                    embedding=embeddings,
                                    url=qdrant_url,
                                    api_key=qdrant_api_key,
                                    collection_name="knowledge"
                                )
                            else:
                                vectorstore = QdrantVectorStore.from_documents(
                                    documents=splits,
                                    embedding=embeddings,
                                    url=qdrant_url,
                                    collection_name="knowledge"
                                )
                    
                    st.session_state.vectorstore = vectorstore
                else:
                    st.warning("没有成功解析任何文档")
        else:
            st.warning("请先上传文档")
    
    # 显示知识库状态
    st.divider()
    try:
        client = get_qdrant_client()
        collections = client.get_collections().collections
        collection_names = [c.name for c in collections]
        if "knowledge" in collection_names:
            st.info("📊 知识库状态：已构建")
        else:
            st.info("📊 知识库状态：未构建")
    except:
        st.info("📊 知识库状态：未连接")
    
    # ============================================
    # 检索参数设置
    # ============================================
    st.divider()
    st.header("⚙️ 检索参数")
    
    # 检索模式
    search_mode = st.radio(
        "检索模式",
        ["相似度优先", "阈值过滤"],
        help="相似度优先：返回最相关的k个文档\n阈值过滤：只返回相似度超过阈值的结果"
    )
    
    # 检索数量
    top_k = st.slider(
        "返回文档数量 (k)",
        min_value=1,
        max_value=5,
        value=2,
        help="最多返回多少个相关文档"
    )
    
    # 阈值设置（仅阈值过滤模式）
    if search_mode == "阈值过滤":
        score_threshold = st.slider(
            "距离阈值",
            min_value=0.0,
            max_value=2.0,
            value=1.0,
            step=0.1,
            help="距离越小越相似\n阈值越小越严格（只返回距离小于阈值的结果）"
        )
        st.caption(f"当前阈值：{score_threshold}（距离越小越相似，阈值越小越严格）")
        st.caption("💡 推荐：0.5-1.5 之间，根据实际效果调整")
    else:
        score_threshold = None
    
    # 保存到 session_state
    st.session_state.search_mode = search_mode
    st.session_state.top_k = top_k
    st.session_state.score_threshold = score_threshold
    
    # ============================================
    # 回答模式设置
    # ============================================
    st.divider()
    st.header("🤖 回答模式")
    
    answer_mode = st.radio(
        "回答策略",
        ["知识库优先", "混合模式"],
        help="知识库优先：只回答知识库中的内容，没有则拒绝\n混合模式：知识库优先，没有则用LLM自己的知识回答"
    )
    
    if answer_mode == "知识库优先":
        st.caption("🔒 严格模式：答案必须有据可查")
    else:
        st.caption("🌐 混合模式：知识库+LLM通用知识")
    
    st.session_state.answer_mode = answer_mode

# ============================================
# 主区域：问答
# ============================================
st.header("💬 智能问答")

# 初始化
if "messages" not in st.session_state:
    st.session_state.messages = []

# 显示历史消息
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# 用户输入
if prompt := st.chat_input("输入你的问题..."):
    # 检查知识库
    try:
        client = get_qdrant_client()
        collections = client.get_collections().collections
        collection_names = [c.name for c in collections]
        if "knowledge" not in collection_names:
            st.error("请先上传文档构建知识库")
    except Exception as e:
        st.error(f"连接Qdrant失败: {e}")
    else:
        # 显示用户消息
        with st.chat_message("user"):
            st.markdown(prompt)
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        # 生成回答
        with st.chat_message("assistant"):
            with st.spinner("思考中..."):
                # 加载向量数据库
                embeddings = BailianEmbeddings()
                client = get_qdrant_client()
                
                vectorstore = QdrantVectorStore(
                    client=client,
                    collection_name="knowledge",
                    embedding=embeddings
                )
                
                # 创建 RAG 链
                llm = ChatOpenAI(
                    model="qwen-turbo",
                    openai_api_key=os.getenv("DASHSCOPE_API_KEY"),
                    openai_api_base="https://dashscope.aliyuncs.com/compatible-mode/v1"
                )
                
                # 使用侧边栏参数创建检索器
                search_mode = st.session_state.get("search_mode", "相似度优先")
                top_k = st.session_state.get("top_k", 2)
                score_threshold = st.session_state.get("score_threshold", None)
                
                if search_mode == "阈值过滤" and score_threshold is not None:
                    retriever = vectorstore.as_retriever(
                        search_type="similarity_score_threshold",
                        search_kwargs={"k": top_k, "score_threshold": score_threshold}
                    )
                else:
                    retriever = vectorstore.as_retriever(search_kwargs={"k": top_k})
                
                # 构建对话历史
                def get_chat_history():
                    """获取对话历史"""
                    history = []
                    for msg in st.session_state.messages:
                        if msg["role"] == "user":
                            history.append(("user", msg["content"]))
                        else:
                            history.append(("assistant", msg["content"]))
                    return history
                
                # 创建带上下文的 prompt
                rag_prompt = ChatPromptTemplate.from_messages([
                    ("system", """你是企业知识库助手，根据以下文档内容回答用户问题。
如果文档中没有相关信息，你可以按照自己的知识回答，此时需要标注结果仅供参考。

文档内容：
{context}"""),
                    MessagesPlaceholder(variable_name="chat_history"),
                    ("user", "{question}")
                ])
                
                def format_docs(docs):
                    if not docs:
                        return ""
                    return "\n\n".join([f"【来源: {doc.metadata.get('source', '未知')}】\n{doc.page_content}" for doc in docs])
                
                # 创建简单的链
                chain = rag_prompt | llm | StrOutputParser()
                
                # 检索文档（带分数）
                docs_with_scores = vectorstore.similarity_search_with_score(prompt, k=top_k)
                
                # 根据阈值过滤
                if search_mode == "阈值过滤" and score_threshold is not None:
                    docs = [doc for doc, score in docs_with_scores if score <= score_threshold]
                else:
                    docs = [doc for doc, score in docs_with_scores]
                
                # 获取对话历史
                chat_history = get_chat_history()
                print(f"his对话:{chat_history}")
                if docs:
                    # 显示参考文档
                    with st.expander("📄 参考文档"):
                        for i, (doc, score) in enumerate([(d, s) for d, s in docs_with_scores if d in docs]):
                            # score 是距离，越小越相似
                            st.text(f"文档 {i+1}: {doc.metadata.get('source', '未知')} (距离: {score:.4f})")
                            st.caption(f"距离越小越相似，当前距离 {score:.4f}")
                            st.text(doc.page_content[:200] + "...")
                    
                    # 生成答案
                    response = chain.invoke({
                        "context": format_docs(docs),
                        "question": prompt,
                        "chat_history":chat_history
                    })
                    st.markdown(response)
                    # 保存回答到消息历史
                    st.session_state.messages.append({"role": "assistant", "content": response})
                else:
                    if answer_mode=="知识库优先":
                       st.markdown("知识库中没有相关信息")
                       st.session_state.messages.append({"role": "assistant", "content": "知识库中没有相关信息"})
                    else:
                         # 知识库没有找到，让 LLM 用自己的知识回答
                        st.warning("知识库中没有找到相关信息，以下回答来自 LLM 通用知识：")
                        response = chain.invoke({
                            "context": "",
                            "question": prompt
                        })
                        st.markdown(response)
                        if search_mode == "阈值过滤":
                            st.info(f"💡 提示：当前距离阈值 {score_threshold}，可尝试提高阈值获取更多结果（距离越小越相似）")
                        st.session_state.messages.append({"role": "assistant", "content": response})

# ============================================
# 底部说明
# ============================================
st.divider()
st.caption("💡 提示：支持上传 TXT、PDF、DOCX、Excel 格式的文档，点击【构建知识库】，然后开始问答")
